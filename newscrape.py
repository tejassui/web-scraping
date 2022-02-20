import requests
from urllib.request import urlopen as ureq #naming urlopen function from library urllib.request as ureq
from bs4 import BeautifulSoup as soup 
import win32com.client as wincl
spk=wincl.Dispatch("SAPI.SpVoice")
from docx import Document
from docx.shared import Pt
import datetime

url = "https://www.indiatoday.in/top-stories"

client = ureq(url)
client  #stores all data from url link in client

page_html = client.read()  #whole website data(backend code) will be stored in page_html
client.close()
page_html
page_soup = soup(page_html,"html.parser")  #parser is used to transfer html language into python readable language
page_soup


articles = page_soup.find("div",{"class":"view-content"}).findAll("div",{"class":"catagory-listing"})
len(articles)

#converting to word document
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = "Calibri"
font.size = Pt(12)
#font.italic = True 
#font.underline = True 

doc.add_heading("Top Stories",0) #0 is for size of heading
doc.add_heading("SOURCE : India Today", 1)
doc.add_heading("NUMBER OF ARTICLES : " + str(len(articles)), 4)


i=1
for x in articles:
    title=x.find("h2").text
    para=x.find("p").text
    #print(i,title)
    #spk.Speak(title)
    #print("\n",para)
    #spk.Speak(para)
    doc.add_paragraph(str(i) + ". " + title)
    doc.add_paragraph(para + "\n")
    i=i+1
    
filename = "news_" + str(datetime.datetime.now()).replace(":",".").replace(" ", "_")[:16] + ".docx"    
doc.save("C:/Users/dell/Desktop" + filename)
print("File Saved")

a = str(datetime.datetime.now())
a[:16]
